import os
import sys
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

spinner = ["/", "-", "\\", "|"] # Define the spinner frames

print("\n==================================")
print(".docx file formatting cleaner v0.5")
print("==================================")
print("\nStrips formatting from docx files resetting the selected styles to defaults.\nPreserves: italic, underline, bold, strikeout. \nOptionally strips metadata.\nOptionally identifies isolated characters formatted differently from surrounding characters/punctuation and extraneous carriage returns.")

def set_run_language(run, lang_code):
    rPr = run._element.get_or_add_rPr()
    for attr in ['w:val', 'w:eastAsia', 'w:bidi']:
        lang = OxmlElement('w:lang')
        lang.set(qn(attr), lang_code)
        rPr.append(lang)

def review_isolated_formatting(doc):
    print("\n--- Starting Isolated Formatting Review ---")
    for para in doc.paragraphs:
        # We need to track the character position manually
        current_pos = 0
        full_text = para.text

        for run in para.runs:
            # Show that we're doing something - update the spinner on the current line
            # \r moves the cursor to the start of the line, end="" prevents a newline
            sys.stdout.write(f"\r {spinner[current_pos % len(spinner)]} Analyzing.")
            sys.stdout.flush()
            clean_text = run.text.strip()
            run_len = len(run.text)

            # Check for isolated formatting (1 character only)
            if len(clean_text) == 1:
                active_formats = []
                if run.bold: active_formats.append("Bold")
                if run.italic: active_formats.append("Italic")
                if run.underline: active_formats.append("Underline")
                if run.font.strike: active_formats.append("Strikethrough")

                if active_formats:
                    # Create a window: 30 chars before, 30 chars after
                    start = max(0, current_pos - 30)
                    end = min(len(full_text), current_pos + 30)

                    # Highlight the specific character in the context string
                    before = full_text[start:current_pos]
                    after = full_text[current_pos + 1:end]
                    # We wrap the character in [[ ]] so you can see it if it's a space
                    window = f"{before}[[{run.text}]]{after}"

                    print(f"\nContext: ...{window}...")
                    print(f"Target: '{run.text}' | Formatting: [{', '.join(active_formats)}]")

                    choice = input("Keep formatting? [y]es / [n]o (revert to plain): ").lower()

                    if choice == 'n':
                        run.bold = run.italic = run.underline = run.font.strike = False
                        print("Reverted.")

            current_pos += run_len

def ultimate_clean_docx():
    if len(sys.argv) > 2:
        input_file, output_file = sys.argv[1], sys.argv[2]
    else:
        input_file = input("\nEnter input .docx path: ").strip('"')
        output_file = input("Enter output .docx path: ").strip('"')

    if not os.path.exists(input_file):
        print("Error: File not found.")
        return

    doc = Document(input_file)
    lang_code = input("\nEnter language code (e.g., en-US) or Enter to skip: ").strip()

    # 1. Map styles
    used_styles = {p.style.name for p in doc.paragraphs}
    style_map = {}
    for name in sorted(used_styles):
        print(f"Style: '{name}'")
        choice = input("  1: Normal, 2: Heading 1, 3: Heading 2, [Enter]: Skip: ")
        if choice == '1': style_map[name] = 'Normal'
        elif choice == '2': style_map[name] = 'Heading 1'
        elif choice == '3': style_map[name] = 'Heading 2'

    # 2. Process Paragraphs
    spin_count = 0
    for para in doc.paragraphs:
        spin_count = spin_count + 1
        print("Processing paragraph styles...")
        # Show that we're doing something - update the spinner on the current line
        # \r moves the cursor to the start of the line, end="" prevents a newline
        sys.stdout.write(f"\r {spinner[spin_count % len(spinner)]} Analyzing.")
        sys.stdout.flush()
        if para.style.name in style_map:
            target_style = style_map[para.style.name]
            # Only try to apply if it exists in the document's style gallery
            if target_style in existing_style_names:
                para.style = doc.styles[target_style]
            else:
                # If it doesn't exist, we just leave it alone rather than crashing
                pass

        # Reset Paragraph geometry
        pf = para.paragraph_format
        pf.line_spacing = pf.space_before = pf.space_after = pf.alignment = None
        pf.left_indent = pf.right_indent = pf.first_line_indent = None

        for run in para.runs:
            # Skip hidden text entirely
            if run.font.hidden:
                run.text = ""
                continue

            run.style = None
            b, i, u, s = run.bold, run.italic, run.underline, run.font.strike

            rPr = run._element.get_or_add_rPr()
            tags_to_kill = [
                qn('w:rFonts'), qn('w:sz'), qn('w:szCs'),
                qn('w:color'), qn('w:highlight'), # Removes Highlighting
                qn('w:shd'),                      # Removes Paragraph/Text Shading
                qn('w:u'),                        # We reset underline via python-docx
                qn('w:ascii'), qn('w:hAnsi'), qn('w:cs')
            ]

            for tag in tags_to_kill:
                element = rPr.find(tag)
                if element is not None:
                    rPr.remove(element)

            run.bold, run.italic, run.underline, run.font.strike = b, i, u, s
            if lang_code:
                set_run_language(run, lang_code)

    # 3 Remove Empty Paragraphs
    print("\nRemove empty paragraphs?")
    rem_choice = input("  1: YES, [Enter]: Skip: ")
    if rem_choice == '1':
        # We use list() to create a static list because we are 
        # modifying the document structure while iterating
        spin_count = 0
        for para in list(doc.paragraphs):
            spin_count = spin_count + 1
            print("Processing paragraph styles...")
            # Show that we're doing something - update the spinner on the current line
            # \r moves the cursor to the start of the line, end="" prevents a newline
            sys.stdout.write(f"\r {spinner[spin_count % len(spinner)]} Analyzing.")
            sys.stdout.flush()
            if not para.text.strip():
                p = para._element
                p.getparent().remove(p)
                p._p = p._element = None
    
    # 4 Review isolated characters
    print("\nWould you like to review isolated formatted characters (e.g., single bold letters)?")
    review_choice = input("  1: YES, [Enter]: Skip: ")
    if review_choice == '1':
        review_isolated_formatting(doc)

    # 5 Strip Metadata
    print("\nStrip Metadata?")
    choice = input("  1: YES, [Enter]: Skip: ")
    if choice == '1':
        core_props = doc.core_properties
        core_props.author = ""
        core_props.comments = ""
        core_props.keywords = ""
        core_props.last_modified_by = ""
        core_props.title = ""

    doc.save(output_file)
    print(f"\nDocument fully scrubbed and saved to: {output_file}")

if __name__ == "__main__":
    ultimate_clean_docx()
